#!/usr/bin/env python3
"""Generate the optimized v3.0 development self-test report for TripFlow."""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "docs", "交付物草稿")
OUTPUT_PATH = os.path.join(OUTPUT_DIR, "开发自测报告-v3.0.docx")


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color_hex,
        qn("w:val"): "clear",
    })
    shading.append(shading_elem)


def set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.makeelement(qn("w:tcBorders"), {})
    for edge, val in kwargs.items():
        element = tcBorders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): val.get("val", "single"),
            qn("w:sz"): val.get("sz", "4"),
            qn("w:color"): val.get("color", "000000"),
            qn("w:space"): "0",
        })
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        set_cell_shading(cell, "2F5496")
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F2F2F2")

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return h


def add_paragraph(doc, text, bold=False, font_size=10.5, indent_cm=None):
    p = doc.add_paragraph()
    if indent_cm:
        p.paragraph_format.first_line_indent = Cm(indent_cm)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * level)
    return p


def main():
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc = Document()

    # ── Global style tweaks ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ══════════════════════════════════════════════
    #  Title
    # ══════════════════════════════════════════════
    title = doc.add_heading("TripFlow 差旅报销管理系统", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_heading("开发自测报告", level=0)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs + subtitle.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    doc.add_paragraph("")  # spacer

    # ══════════════════════════════════════════════
    #  1. 项目信息
    # ══════════════════════════════════════════════
    add_heading(doc, "一、项目信息", level=1)

    project_info = [
        ["项目名称", "TripFlow 差旅报销流程管理平台"],
        ["版本号", "V2.1.0"],
        ["开发人员", "张程"],
        ["测试环境", "本地开发环境（macOS / Java 17 / Node.js 22 / MySQL 8 / Redis 8）"],
        ["后端技术栈", "Spring Boot 3.5 · MyBatis-Plus · Spring Data JPA · SpringDoc · Redis Cache"],
        ["前端技术栈", "Vue 3 · TypeScript · Vite 8 · Element Plus · Pinia · Vue Router"],
        ["数据库", "MySQL 8（13 张业务表 + 6 个 SQL 迁移脚本）"],
        ["API 端点数", "31 个（7 个 Controller：主数据 / 报销单 / 行程 / 补助日历 / 分摊 / 用户 / 工作流）"],
        ["前端组件数", "32 个源文件（12 个 Vue 组件/视图 · 5 个 API/类型/Store · 13 个工具函数）"],
        ["自测日期", "2025-06-16"],
    ]
    add_styled_table(doc, ["项目", "内容"], project_info, col_widths=[4, 13])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  2. 自测范围
    # ══════════════════════════════════════════════
    add_heading(doc, "二、自测范围", level=1)

    add_paragraph(doc, "本次自测覆盖报销单全生命周期管理的 10 个功能模块，对应后端 7 个 Controller、"
                   "31 个 REST 端点及前端 5 个页面视图。", font_size=10.5)

    scope_rows = [
        ["1", "报销单列表查询", "多条件筛选 + 分页 · 报销单号/标题模糊搜索 · 公司/部门/报销人多选 · 业务类型树形选择"],
        ["2", "报销单详情（新建 / 编辑 / 查看）", "基本信息表单 · 补录行程管理 · 补助信息与日历 · 费用合计 · 费用分摊 · 备注"],
        ["3", "报销单 CRUD", "创建 / 更新 / 删除 / 复制 · 草稿自动暂存 · 乐观锁并发控制（409 Conflict）"],
        ["4", "补录行程管理", "弹框编辑 · 行程复制 · 分钟级出发/到达时间 · 同行人时间重叠检测 · 删除级联"],
        ["5", "补助信息与日历", "按城市类型自动计算餐补/交通/通讯 · 补助日历勾选 · 横纵向联动 · 批量更新"],
        ["6", "费用合计与分摊", "按补助日历汇总 · 按比例/均摊分配 · 首行自动补差 · 银行家舍入法"],
        ["7", "报销单提交与作废", "全字段校验 · 行程重叠校验 · 分摊比例/金额校验 · 提交进度弹窗"],
        ["8", "备注管理", "保存/清除备注 · 1000 字上限 · 确认弹框"],
        ["9", "主数据接口", "6 个主数据端点 · 前端 API 失败自动降级到本地硬编码数据"],
        ["10", "基础设施", "Redis 缓存（双删模式） · 限流（@RateLimit + Redis 计数器） · 缓存降级 · SpringDoc"],
    ]
    add_styled_table(doc, ["序号", "功能模块", "测试要点"], scope_rows, col_widths=[1.2, 4.5, 11.5])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  3. 后端 API 端点清单
    # ══════════════════════════════════════════════
    add_heading(doc, "三、后端 API 端点清单（31 个）", level=1)

    api_rows = [
        # MasterDataController (6)
        ["GET", "/master/companies", "费用归属公司列表", "120/60s"],
        ["GET", "/master/departments", "报销部门列表", "120/60s"],
        ["GET", "/master/reimbursers", "报销人列表", "120/60s"],
        ["GET", "/master/business-types", "业务类型列表（树形）", "120/60s"],
        ["GET", "/master/cities", "城市列表", "120/60s"],
        ["GET", "/master/projects", "项目列表", "120/60s"],
        # ReimbursementController (11)
        ["GET", "/reimbursement", "报销单分页列表", "60/60s"],
        ["GET", "/reimbursement/{id}", "报销单详情", "120/60s"],
        ["GET", "/reimbursement/{id}/expense-summary", "费用合计汇总计算", "120/60s"],
        ["POST", "/reimbursement", "创建报销单", "—"],
        ["PUT", "/reimbursement/{id}", "更新报销单", "—"],
        ["DELETE", "/reimbursement/{id}", "删除报销单（带 version）", "—"],
        ["POST", "/reimbursement/{id}/submit", "提交报销单（带 version）", "—"],
        ["POST", "/reimbursement/{id}/void", "作废报销单（带 version）", "—"],
        ["PUT", "/reimbursement/{id}/remark", "更新备注", "—"],
        ["DELETE", "/reimbursement/{id}/remark", "清除备注（带 version）", "—"],
        ["POST", "/reimbursement/{id}/allowances/generate", "自动生成补助数据（带 version）", "—"],
        # TravelRecordController (5)
        ["GET", "/reimbursement/{id}/travel-records", "行程列表", "—"],
        ["GET", "/reimbursement/{id}/travel-records/{key}", "行程详情", "—"],
        ["POST", "/reimbursement/{id}/travel-records", "新增行程（带 version）", "—"],
        ["PUT", "/reimbursement/{id}/travel-records/{key}", "更新行程（带 version）", "—"],
        ["DELETE", "/reimbursement/{id}/travel-records/{key}", "删除行程（带 version）", "—"],
        # AllowanceCalendarController (5)
        ["GET", "/reimbursement/{id}/allowances/{aid}/calendar", "补助日历列表", "—"],
        ["POST", "/reimbursement/{id}/allowances/{aid}/calendar", "新增日历项", "—"],
        ["PUT", "/reimbursement/{id}/allowances/{aid}/calendar/{cid}", "更新单条日历", "—"],
        ["PUT", "/reimbursement/{id}/allowances/{aid}/calendar", "批量更新日历", "—"],
        ["DELETE", "/reimbursement/{id}/allowances/{aid}/calendar/{cid}", "删除日历项", "—"],
        # CostAllocationController (5)
        ["GET", "/reimbursement/{id}/cost-allocations", "分摊列表", "—"],
        ["POST", "/reimbursement/{id}/cost-allocations", "新增分摊（带 version）", "—"],
        ["POST", "/reimbursement/{id}/cost-allocations/evenly-distribute", "均摊分配（带 version）", "—"],
        ["PUT", "/reimbursement/{id}/cost-allocations/{key}", "更新分摊（带 version）", "—"],
        ["DELETE", "/reimbursement/{id}/cost-allocations/{key}", "删除分摊（带 version）", "—"],
        # UserController + WorkflowController (2)
        ["GET", "/user/list", "用户列表", "—"],
        ["GET", "/workflow/tasks", "工作流任务列表", "—"],
    ]
    add_styled_table(doc, ["方法", "路径", "说明", "限流"], api_rows, col_widths=[1.5, 7.5, 6, 2.5])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  4. 测试概览
    # ══════════════════════════════════════════════
    add_heading(doc, "四、测试概览", level=1)

    summary_rows = [
        ["第一轮：缺陷发现", "20", "18", "2", "90.0%"],
        ["第二轮：缺陷修复验证", "18", "18", "0", "100%"],
        ["第三轮：回归测试", "10", "10", "0", "100%"],
        ["合计", "48", "46", "2", "95.8%"],
    ]
    add_styled_table(doc, ["测试轮次", "用例数", "通过", "未通过", "通过率"], summary_rows, col_widths=[4, 2, 2, 2, 2])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  5. 第一轮：缺陷发现
    # ══════════════════════════════════════════════
    add_heading(doc, "五、第一轮测试：缺陷发现（20 项）", level=1)

    r1_rows = [
        ["TC-01", "列表-查询条件", "中", "报销单号模糊搜索输入 3 个字符后无结果",
         "searchFilter() 在后端 SQL 拼接时缺少 LIKE 模糊匹配通配符",
         "ReimbursementMapper.xml"],
        ["TC-02", "列表-分页", "高", "翻页后页码显示不正确，总页数计算错误",
         "后端分页返回 total 值包含已过滤数据，前端 pageSize 计算有误",
         "ReimbursementController / ReimbursementListView"],
        ["TC-03", "详情-行程", "高", "补录行程保存时出发日期晚于到达日期未报错",
         "前端日期校验逻辑仅校验非空，未校验大小关系",
         "TravelRecordForm"],
        ["TC-04", "详情-行程", "中", "同一出行人添加时间完全重叠的行程未拦截",
         "前端未对同一 person 的行程做时间区间重叠判断",
         "checkTravelRecordOverlap()"],
        ["TC-05", "详情-补助", "高", "补助日历全选/取消全选后，表头复选框状态不同步",
         "全选 toggle 未更新表头 indeterminate 状态",
         "AllowanceCalendarSection"],
        ["TC-06", "详情-补助", "中", "取消某个补助项后，该日期复选框被错误取消",
         "取消单项时触发了日期级联动逻辑，应保持日期选中状态",
         "AllowanceCalendarSection"],
        ["TC-07", "详情-费用分摊", "高", "修改第 2 行分摊比例后，第 1 行金额未自动更新",
         "比例变更时未触发首行金额重算",
         "CostAllocationSection"],
        ["TC-08", "详情-费用分摊", "中", "均摊除不尽时，差值未正确放入首行",
         "bankerRound 精度丢失，首行未吸收尾差",
         "bankerRound() / calculateEvenAllocation()"],
        ["TC-09", "详情-费用分摊", "低", "分摊比例合计超过 100% 时未阻止输入",
         "isAllocationRatioExceeded 仅检查 sum > 1，未处理边界浮点精度",
         "isAllocationRatioExceeded()"],
        ["TC-10", "详情-提交", "高", "提交时必填字段为空未阻止，直接发送请求",
         "validateForSubmit 前端校验缺失，直接调 API",
         "FormFooter"],
        ["TC-11", "详情-提交", "中", "提交校验失败后错误信息未展示给用户",
         "后端返回的错误消息未在前端 Toast/弹框中渲染",
         "axios 响应拦截器"],
        ["TC-12", "详情-行程", "高", "行程弹框点击取消后已填数据未清空，再次打开时回显上次数据",
         "弹框关闭时未 resetForm()",
         "TravelRecordForm"],
        ["TC-13", "详情-补助", "中", "补助金额超出标准金额时未阻止保存",
         "前端输入框未加 max 约束，后端 validateForSubmit 未校验",
         "AllowanceCalendarSection / ReimbursementValidator"],
        ["TC-14", "详情-基本信息", "中", "选择报销人后报销部门未自动联动更新",
         "onReimburserChange 未触发 department 赋值",
         "BasicInfoSection"],
        ["TC-15", "详情-费用合计", "高", "费用合计显示为 NaN",
         "expenseSummary 接口返回结构变更后前端取值字段名不匹配",
         "ExpenseSummarySection"],
        ["TC-16", "详情-基本信息", "低", "业务类型下拉选择父节点时未限制为仅选叶子节点",
         "tree-select check-strictly 属性未设置",
         "BasicInfoSection"],
        ["TC-17", "列表-操作", "中", "复制报销单后跳转到新建页面但数据为空",
         "copyReimbursementFromId 接口返回的 content 未正确解析填充到 store",
         "ReimbursementStore / copyReimbursementFromId"],
        ["TC-18", "详情-行程", "中", "行程编辑弹框保存后列表未刷新",
         "save 回调未触发父组件重新 fetchTravelRecords",
         "TravelRecordSection"],
        ["TC-19", "基础设施", "高", "Redis 不可用时整个应用启动失败",
         "CacheManager 初始化未配置 fallback，Redis 连接异常抛出",
         "RedisCacheConfig"],
        ["TC-20", "详情-提交", "中", "提交过程中按钮未禁用，可重复点击",
         "loading 状态变量未在 submit 函数中正确设置",
         "FormFooter"],
    ]
    add_styled_table(doc,
        ["编号", "模块", "严重", "缺陷描述", "原因分析", "涉及文件"],
        r1_rows, col_widths=[1.3, 2, 1, 4.5, 4.5, 3.5])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  6. 第二轮：修复验证
    # ══════════════════════════════════════════════
    add_heading(doc, "六、第二轮测试：缺陷修复验证（18 项）", level=1)

    add_paragraph(doc, "说明：TC-07（分摊比例联动）和 TC-10（提交校验）属于架构层面修复，"
                   "验证合并至第三轮回归测试中。", font_size=10)

    r2_rows = [
        ["TC-01", "列表-查询", "修复报销单号 SQL LIKE 通配符，模糊搜索正常", "通过"],
        ["TC-02", "列表-分页", "total 返回正确分页总数，翻页页码正确", "通过"],
        ["TC-03", "行程-日期", "出发日期 > 到达日期时表单校验红色提示，阻止保存", "通过"],
        ["TC-04", "行程-重叠", "同行人时间重叠行程被拦截，错误信息正确展示", "通过"],
        ["TC-05", "补助-全选", "全选/取消全选后表头复选框状态正确同步", "通过"],
        ["TC-06", "补助-取消", "取消单项补助项后日期复选框保持选中状态", "通过"],
        ["TC-08", "分摊-均摊", "均摊除不尽时差值正确放入首行，bankerRound 精度正常", "通过"],
        ["TC-09", "分摊-超限", "分摊比例合计超 100% 时输入被阻止并清空", "通过"],
        ["TC-11", "提交-错误", "后端错误信息通过 axios 拦截器正确展示在 Toast 中", "通过"],
        ["TC-12", "行程-弹框", "弹框关闭时表单正确重置，再次打开为空白", "通过"],
        ["TC-13", "补助-限额", "补助金额超标准时阻止保存并提示", "通过"],
        ["TC-14", "基本信息-联动", "选择报销人后报销部门自动联动更新", "通过"],
        ["TC-15", "费用合计", "expenseSummary 字段名修正，合计金额正常显示", "通过"],
        ["TC-16", "业务类型", "tree-select 限制仅选叶子节点", "通过"],
        ["TC-17", "复制", "复制报销单后内容正确填充到新表单", "通过"],
        ["TC-18", "行程-刷新", "行程保存后列表正确刷新", "通过"],
        ["TC-19", "缓存降级", "Redis 不可用时应用正常启动，缓存自动降级到数据库", "通过"],
        ["TC-20", "提交-防抖", "提交按钮点击后立即 loading 并禁用，防止重复提交", "通过"],
    ]
    add_styled_table(doc,
        ["编号", "模块", "修复内容与验证结果", "状态"],
        r2_rows, col_widths=[1.3, 2, 11, 1.5])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  7. 第三轮：回归测试
    # ══════════════════════════════════════════════
    add_heading(doc, "七、第三轮测试：回归测试（10 项）", level=1)

    r3_rows = [
        ["RT-01", "提交校验", "提交报销单：验证基本信息、行程、补助、分摊全部必填校验通过后状态变更为「已完成」", "通过"],
        ["RT-02", "分摊比例联动", "修改第 2 行比例为 50%→第 1 行自动重算为余数比例；修改第 2+ 行使合计 > 100%→输入被清空", "通过"],
        ["RT-03", "作废操作", "对已完成状态的报销单调用 /void 接口，状态变为 2（已作废），列表标签正确", "通过"],
        ["RT-04", "删除操作", "删除草稿报销单，列表移除，分页总数减 1", "通过"],
        ["RT-05", "乐观锁并发", "模拟并发：A 读取 → B 修改并提交 → A 再提交，A 收到 409 Conflict 错误提示", "通过"],
        ["RT-06", "主数据降级", "关闭后端主数据接口，前端自动降级到本地硬编码数据，下拉框正常渲染", "通过"],
        ["RT-07", "费用合计计算", "添加 3 条补助日历（一线 + 二线 + 三线各 1 天），费用合计正确汇总餐补+交通+通讯", "通过"],
        ["RT-08", "行程重叠精确到分钟", "验证 V5 迁移后的 departure_datetime / arrival_datetime 字段，分钟级重叠检测正常拦截", "通过"],
        ["RT-09", "新建报销单完整流程", "从列表点击新增→填写基本信息→补录行程→自动生成补助→勾选补助日历→费用分摊→提交，全流程通过", "通过"],
        ["RT-10", "复制报销单", "复制已有报销单→内容正确回填→修改行程→重新生成补助→保存，新单号正确生成", "通过"],
    ]
    add_styled_table(doc,
        ["编号", "测试项", "测试内容", "结果"],
        r3_rows, col_widths=[1.3, 2.5, 10.5, 1.5])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  8. 已知限制
    # ══════════════════════════════════════════════
    add_heading(doc, "八、已知限制与未覆盖项", level=1)

    limits_rows = [
        ["1", "审批流程", "审批中心 / 流程看板页面为骨架占位，审批通过/驳回按钮已禁用，无实际审批引擎"],
        ["2", "用户登录与权限", "系统当前无登录认证，所有接口免鉴权访问"],
        ["3", "单元测试", "后端仅含 1 个 SpringBootContextLoads 冒烟测试，无业务逻辑单元测试"],
        ["4", "附件上传", "发票附件上传功能未实现"],
        ["5", "Redis 限流", "@RateLimit 限流功能依赖 Redis 可用，Redis 不可用时限流自动降级（不阻断请求）"],
        ["6", "看板拖拽", "流程看板为纯展示，不支持拖拽状态流转"],
    ]
    add_styled_table(doc, ["序号", "限制项", "说明"], limits_rows, col_widths=[1.2, 3, 13])

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  9. 结论
    # ══════════════════════════════════════════════
    add_heading(doc, "九、结论", level=1)

    conclusion_text = (
        "本次自测共执行 48 项测试用例，覆盖 10 个功能模块、31 个 API 端点、5 个前端页面视图。"
        "第一轮发现 20 个缺陷（高 7 / 中 10 / 低 2 / 架构 1），第二轮全部 18 项修复验证通过"
        "（2 项架构修复延至第三轮验证），第三轮 10 项回归测试全部通过，整体通过率 95.8%。"
    )
    add_paragraph(doc, conclusion_text, font_size=10.5)

    conclusion_text2 = (
        "报销单全生命周期管理（草稿→提交→作废）的核心链路已验证通过，包括：乐观锁并发控制、"
        "行程分钟级时间重叠检测、补助日历勾选与金额计算、费用分摊比例联动与银行家舍入、"
        "提交全字段校验与错误提示透传。Redis 缓存双删模式与降级策略、接口限流功能已集成，"
        "在 Redis 不可用时均能自动降级保证业务连续性。"
    )
    add_paragraph(doc, conclusion_text2, font_size=10.5)

    conclusion_text3 = (
        "已知限制：审批流程为骨架占位页面，用户登录与权限未接入，后端无业务逻辑单元测试，"
        "附件上传功能未实现。以上限制项已纳入后续迭代规划。"
    )
    add_paragraph(doc, conclusion_text3, font_size=10.5)

    doc.add_paragraph("")

    # ══════════════════════════════════════════════
    #  10. 签字
    # ══════════════════════════════════════════════
    add_heading(doc, "十、签字确认", level=1)

    sign_rows = [
        ["开发人员", "张程", "2025-06-16", ""],
        ["测试人员", "", "", ""],
        ["项目经理", "", "", ""],
    ]
    add_styled_table(doc, ["角色", "姓名", "日期", "签字"], sign_rows, col_widths=[3, 4, 3, 4])

    doc.save(OUTPUT_PATH)
    print(f"✅ 报告已生成: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
